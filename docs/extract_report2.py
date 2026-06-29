from docx import Document
import sys

doc = Document(r'D:\cyberaiuz\docs\CyberAI_UZ_Audit_Report_2026.docx')
output = []

for para in doc.paragraphs:
    if para.text.strip():
        style_name = para.style.name if para.style else ''
        prefix = ''
        if 'Heading 1' in style_name:
            prefix = '# '
        elif 'Heading 2' in style_name:
            prefix = '## '
        elif 'Heading 3' in style_name:
            prefix = '### '
        elif 'Heading 4' in style_name:
            prefix = '#### '
        output.append(prefix + para.text)

# Also extract tables
for table in doc.tables:
    output.append('\n--- TABLE ---')
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        output.append(' | '.join(cells))
    output.append('--- END TABLE ---\n')

with open(r'D:\cyberaiuz\docs\audit_report_2.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))

print(f"Extracted {len(output)} lines from Report 2")
